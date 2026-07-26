from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/eeri/coding/monorepo/projects/cleargatecustoms")
OUT = ROOT / "output" / "docx" / "Aleksandr_Hubanov_Resume.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Arial"
INK = RGBColor(26, 35, 50)
ACCENT = RGBColor(28, 78, 121)
MUTED = RGBColor(82, 92, 105)
RULE = "B9C6D3"


def set_cellless_font(run, size, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_bottom_border(paragraph, color=RULE, size="8", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text, url, size=9.2, color=ACCENT, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_el)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_separator(paragraph):
    run = paragraph.add_run("  |  ")
    set_cellless_font(run, 8.7, color=MUTED)


def add_custom_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "330")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "330")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "36")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    r_pr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1C4E79")
    r_pr.append(color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def add_section_heading(document, text):
    p = document.add_paragraph(style="Resume Section")
    p.add_run(text.upper())
    add_bottom_border(p)
    return p


def add_role_header(document, company, role, dates, location, descriptor):
    p = document.add_paragraph(style="Role Header")
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(company)
    set_cellless_font(r, 10.2, bold=True, color=INK)
    r = p.add_run(f"  -  {role}")
    set_cellless_font(r, 10.2, bold=True, color=ACCENT)
    r = p.add_run("\t" + dates)
    set_cellless_font(r, 9.3, bold=True, color=INK)

    meta = document.add_paragraph(style="Role Meta")
    r = meta.add_run(descriptor)
    set_cellless_font(r, 8.7, italic=True, color=MUTED)
    r = meta.add_run("\t" + location)
    set_cellless_font(r, 8.7, italic=True, color=MUTED)
    return p, meta


def add_bullet(document, text, num_id):
    p = document.add_paragraph(style="Resume Bullet")
    apply_num(p, num_id)
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.54)
section.bottom_margin = Inches(0.54)
section.left_margin = Inches(0.70)
section.right_margin = Inches(0.70)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

# Compact reference guide, with named resume-specific one-page geometry/type overrides.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(9.8)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
section_style.font.name = FONT
section_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
section_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
section_style.font.size = Pt(10.5)
section_style.font.bold = True
section_style.font.color.rgb = ACCENT
section_style.paragraph_format.space_before = Pt(7.2)
section_style.paragraph_format.space_after = Pt(3.6)
section_style.paragraph_format.keep_with_next = True

role_header_style = styles.add_style("Role Header", WD_STYLE_TYPE.PARAGRAPH)
role_header_style.paragraph_format.space_before = Pt(1)
role_header_style.paragraph_format.space_after = Pt(0)
role_header_style.paragraph_format.keep_with_next = True

role_meta_style = styles.add_style("Role Meta", WD_STYLE_TYPE.PARAGRAPH)
role_meta_style.paragraph_format.space_before = Pt(0)
role_meta_style.paragraph_format.space_after = Pt(1.4)
role_meta_style.paragraph_format.keep_with_next = True
role_meta_style.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)

bullet_style = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
bullet_style.font.name = FONT
bullet_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
bullet_style.font.size = Pt(9.55)
bullet_style.font.color.rgb = INK
bullet_style.paragraph_format.space_before = Pt(0)
bullet_style.paragraph_format.space_after = Pt(2.3)
bullet_style.paragraph_format.line_spacing = 1.0
bullet_style.paragraph_format.keep_together = True

name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_before = Pt(0)
name.paragraph_format.space_after = Pt(0.5)
r = name.add_run("ALEKSANDR HUBANOV")
set_cellless_font(r, 21.5, bold=True, color=INK)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(2.6)
r = title.add_run("FULL-STACK DEVELOPER")
set_cellless_font(r, 10.9, bold=True, color=ACCENT)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_before = Pt(0)
contact.paragraph_format.space_after = Pt(4)
contact.paragraph_format.line_spacing = 1.0
add_hyperlink(contact, "+380 63 771 51 16", "tel:+380637715116")
add_separator(contact)
add_hyperlink(contact, "eeri.dev@gmail.com", "mailto:eeri.dev@gmail.com")
add_separator(contact)
add_hyperlink(contact, "linkedin.com/in/aleksandr-hubanov", "https://www.linkedin.com/in/aleksandr-hubanov/")
add_separator(contact)
add_hyperlink(contact, "eeri.dev", "https://eeri.dev/")
add_separator(contact)
r = contact.add_run("Romania")
set_cellless_font(r, 9.2, color=MUTED)
add_bottom_border(contact, color="8097AC", size="10", space="4")

summary = doc.add_paragraph()
summary.paragraph_format.space_before = Pt(4)
summary.paragraph_format.space_after = Pt(1.5)
summary.paragraph_format.line_spacing = 1.0
r = summary.add_run(
    "Full-stack developer experienced in building maintainable web applications and REST APIs "
    "with Python, Django/FastAPI, PostgreSQL, React, and TypeScript. Strong focus on backend "
    "architecture, authentication, API integration, reusable frontend systems, and Docker-based environments."
)
set_cellless_font(r, 9.8, color=INK)

bullet_num_id = add_custom_bullet_numbering(doc)

add_section_heading(doc, "Experience")
add_role_header(
    doc,
    "IVE STUDIO",
    "Full-Stack Developer",
    "Feb 2025 - Apr 2026",
    "Remote",
    "Product development studio focused on web platforms and digital products",
)
for text in [
    "Designed and delivered REST APIs with Python, Django, and PostgreSQL, supporting reliable data exchange across web platform workflows.",
    "Implemented authentication and role-based access control, centralizing permission enforcement for platform users.",
    "Built React and TypeScript interfaces connected to backend services, with reusable components and a scalable frontend structure.",
    "Standardized the development environment with Docker and Docker Compose, improving consistency across local and deployment workflows.",
]:
    add_bullet(doc, text, bullet_num_id)

add_role_header(
    doc,
    "AI-DEF",
    "Full-Stack Developer",
    "Apr 2023 - Dec 2024",
    "Remote",
    "Defense technology company developing autonomous systems and UAV platforms",
)
for text in [
    "Developed Django backend services and APIs for internal platforms supporting data management and operational workflows.",
    "Modeled, queried, and maintained application data in PostgreSQL through the Django ORM.",
    "Integrated frontend interfaces with backend APIs, enabling web-based access to internal tools and platform workflows.",
    "Containerized development and deployment environments with Docker to improve reproducibility across the delivery process.",
]:
    add_bullet(doc, text, bullet_num_id)

add_section_heading(doc, "Technical Skills")
skills = [
    ("Languages", "Python, TypeScript, HTML, CSS"),
    ("Backend", "Django, FastAPI, REST APIs, authentication, role-based access control"),
    ("Frontend", "React, reusable component architecture, API integration"),
    ("Data & Infrastructure", "PostgreSQL, Django ORM, Docker, Docker Compose"),
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.35)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(label + ": ")
    set_cellless_font(r, 9.45, bold=True, color=INK)
    r = p.add_run(value)
    set_cellless_font(r, 9.45, color=INK)

add_section_heading(doc, "Languages")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0
for idx, (label, value) in enumerate(
    [("English", "B2"), ("Russian", "Native"), ("Ukrainian", "Native")]
):
    if idx:
        add_separator(p)
    r = p.add_run(label + ": ")
    set_cellless_font(r, 9.45, bold=True, color=INK)
    r = p.add_run(value)
    set_cellless_font(r, 9.45, color=INK)

doc.core_properties.title = "Aleksandr Hubanov - Full-Stack Developer Resume"
doc.core_properties.subject = "Universal one-page resume"
doc.core_properties.author = "Aleksandr Hubanov"
doc.core_properties.keywords = (
    "Full-Stack Developer, Python, Django, FastAPI, PostgreSQL, React, TypeScript, Docker"
)
doc.save(OUT)
print(OUT)
