from copy import deepcopy
from hashlib import sha256
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REFERENCE = Path("/Users/eeri/Downloads/Копия ChaseTalks Resume Template.docx")
EXPECTED_SHA = "4982166ecdad18a949395bdc4c46ff2e03be3857ace7607a7d9bebcd0c0228a5"
OUTPUT = Path(
    "/Users/eeri/coding/monorepo/projects/cleargatecustoms/output/docx/"
    "Aleksandr_Hubanov_Resume_ChaseTalks_Style.docx"
)
WORKING = Path(
    "/Users/eeri/coding/monorepo/projects/cleargatecustoms/tmp/chase_template/"
    "working_generated.docx"
)


def file_sha(path):
    h = sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def clear_paragraph_contents(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_run_font(run, name="Rubik", size=9, bold=None, color="282828"):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Rubik")
    r_fonts.set(qn("w:hAnsi"), "Rubik")
    r_fonts.set(qn("w:cs"), "Rubik")
    r_pr.append(r_fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "282828")
    r_pr.append(color)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_contact_separator(paragraph):
    run = paragraph.add_run("  -  ")
    set_run_font(run, size=10)


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def paragraph_from_template(cell, template_paragraph):
    cloned = deepcopy(template_paragraph._p)
    cell._tc.append(cloned)
    return cell.paragraphs[-1]


def set_simple_text(paragraph, text):
    texts = paragraph._p.findall(".//" + qn("w:t"))
    if not texts:
        run = paragraph.add_run(text)
        set_run_font(run)
        return
    texts[0].text = text
    if text.startswith(" ") or text.endswith(" "):
        texts[0].set(qn("xml:space"), "preserve")
    for extra in texts[1:]:
        extra.text = ""


def fill_heading(row, label, heading_left_template, heading_right_template):
    left = row.cells[0]
    right = row.cells[1]
    clear_cell(left)
    if right._tc is not left._tc:
        clear_cell(right)
    left_p = paragraph_from_template(left, heading_left_template)
    set_simple_text(left_p, label)
    if right._tc is not left._tc:
        right_p = paragraph_from_template(right, heading_right_template)
        set_simple_text(right_p, "")


def fill_split_row(row, left_text, right_text, left_template, right_template):
    left = row.cells[0]
    right = row.cells[1]
    clear_cell(left)
    clear_cell(right)
    left_p = paragraph_from_template(left, left_template)
    right_p = paragraph_from_template(right, right_template)
    set_simple_text(left_p, left_text)
    set_simple_text(right_p, right_text)


def fill_profile_content(cell, template, text):
    clear_cell(cell)
    p = paragraph_from_template(cell, template)
    set_simple_text(p, text)


def fill_experience_content(cell, templates, role, location, bullets):
    clear_cell(cell)
    role_p = paragraph_from_template(cell, templates["regular"])
    set_simple_text(role_p, role)
    location_p = paragraph_from_template(cell, templates["location"])
    set_simple_text(location_p, location)
    paragraph_from_template(cell, templates["blank"])
    for bullet in bullets:
        p = paragraph_from_template(cell, templates["bullet"])
        set_simple_text(p, bullet)
    paragraph_from_template(cell, templates["tail_blank"])


def fill_labeled_lines(cell, blank_template, label_template, lines):
    clear_cell(cell)
    paragraph_from_template(cell, blank_template)
    for label, value in lines:
        p = paragraph_from_template(cell, label_template)
        texts = p._p.findall(".//" + qn("w:t"))
        texts[0].text = label
        texts[0].set(qn("xml:space"), "preserve")
        texts[1].text = value
        for extra in texts[2:]:
            extra.text = ""


if file_sha(REFERENCE) != EXPECTED_SHA:
    raise RuntimeError("Reference template changed; fresh distillation is required.")

doc = Document(REFERENCE)

# Capture source components before editing them.
table = doc.tables[0]
templates = {
    "heading_left": deepcopy(table.rows[0].cells[0].paragraphs[0]),
    "heading_right": deepcopy(table.rows[0].cells[1].paragraphs[0]),
    "company": deepcopy(table.rows[4].cells[0].paragraphs[1]),
    "date": deepcopy(table.rows[4].cells[1].paragraphs[1]),
    "regular": deepcopy(table.rows[5].cells[0].paragraphs[0]),
    "location": deepcopy(table.rows[5].cells[0].paragraphs[1]),
    "blank": deepcopy(table.rows[5].cells[0].paragraphs[2]),
    "bullet": deepcopy(table.rows[5].cells[0].paragraphs[3]),
    "tail_blank": deepcopy(table.rows[5].cells[0].paragraphs[7]),
    "skills_blank": deepcopy(table.rows[11].cells[0].paragraphs[0]),
    "skills_line": deepcopy(table.rows[11].cells[0].paragraphs[1]),
}

# Remove the instructional first page, including its explicit page break.
for paragraph in list(doc.paragraphs[:15]):
    remove_paragraph(paragraph)

# Update masthead while retaining the source paragraph formatting.
body_paragraphs = doc.paragraphs
set_simple_text(body_paragraphs[0], "Aleksandr Hubanov")

contact = body_paragraphs[2]
clear_paragraph_contents(contact)
add_hyperlink(contact, "+380 63 771 51 16", "tel:+380637715116")
add_contact_separator(contact)
add_hyperlink(contact, "eeri.dev@gmail.com", "mailto:eeri.dev@gmail.com")
add_contact_separator(contact)
add_hyperlink(
    contact,
    "linkedin.com/in/aleksandr-hubanov",
    "https://www.linkedin.com/in/aleksandr-hubanov/",
)
add_contact_separator(contact)
add_hyperlink(contact, "eeri.dev", "https://eeri.dev/")

# PROFILE
fill_heading(
    table.rows[0],
    "PROFILE",
    templates["heading_left"],
    templates["heading_right"],
)
fill_split_row(
    table.rows[1],
    "Full-Stack Developer",
    "Romania",
    templates["company"],
    templates["date"],
)
fill_profile_content(
    table.rows[2].cells[0],
    templates["regular"],
    "Full-stack developer building maintainable web applications and REST APIs "
    "with Python, Django/FastAPI, PostgreSQL, React, and TypeScript. Experienced "
    "in backend architecture, authentication, API integrations, reusable frontend "
    "components, and Docker-based environments.",
)

# EXPERIENCE
fill_heading(
    table.rows[3],
    "EXPERIENCE",
    templates["heading_left"],
    templates["heading_right"],
)
fill_split_row(
    table.rows[4],
    "IVE STUDIO",
    "February 2025 - April 2026",
    templates["company"],
    templates["date"],
)
fill_experience_content(
    table.rows[5].cells[0],
    templates,
    "Full-Stack Developer",
    "Remote",
    [
        "Built REST APIs with Python, Django, and PostgreSQL for reliable data exchange across web platform workflows.",
        "Implemented authentication and role-based access control, centralizing permissions for platform users.",
        "Built React/TypeScript interfaces with reusable components and scalable frontend architecture.",
        "Standardized development with Docker Compose, improving consistency across local and deployment workflows.",
    ],
)
fill_split_row(
    table.rows[6],
    "AI-DEF",
    "April 2023 - December 2024",
    templates["company"],
    templates["date"],
)
fill_experience_content(
    table.rows[7].cells[0],
    templates,
    "Full-Stack Developer",
    "Remote",
    [
        "Developed Django services and APIs for internal data management and operational workflows.",
        "Modeled and maintained PostgreSQL application data through the Django ORM.",
        "Connected frontend interfaces to backend APIs for web-based access to internal platform tools.",
        "Containerized development and deployment with Docker to improve delivery reproducibility.",
    ],
)

# TECHNICAL SKILLS
fill_heading(
    table.rows[8],
    "TECHNICAL SKILLS",
    templates["heading_left"],
    templates["heading_right"],
)
fill_labeled_lines(
    table.rows[9].cells[0],
    templates["skills_blank"],
    templates["skills_line"],
    [
        ("Programming Languages: ", "Python, TypeScript, HTML, CSS"),
        ("Backend: ", "Django, FastAPI, REST APIs, authentication, role-based access control"),
        ("Frontend: ", "React, reusable component architecture, API integration"),
        ("Data & Infrastructure: ", "PostgreSQL, Django ORM, Docker, Docker Compose"),
    ],
)

# LANGUAGES
fill_heading(
    table.rows[10],
    "LANGUAGES",
    templates["heading_left"],
    templates["heading_right"],
)
fill_labeled_lines(
    table.rows[11].cells[0],
    templates["skills_blank"],
    templates["skills_line"],
    [("Languages: ", "English (B2), Russian (Native), Ukrainian (Native)")],
)

doc.core_properties.title = "Aleksandr Hubanov - Full-Stack Developer Resume"
doc.core_properties.subject = "Universal one-page resume"
doc.core_properties.author = "Aleksandr Hubanov"
doc.core_properties.keywords = (
    "Full-Stack Developer, Python, Django, FastAPI, PostgreSQL, React, TypeScript, Docker"
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
WORKING.parent.mkdir(parents=True, exist_ok=True)
doc.save(WORKING)

# Preserve every source package part byte-for-byte except the two editable
# parts that hold body content and its external hyperlink relationships.
replacement_parts = {
    "word/document.xml",
    "word/_rels/document.xml.rels",
}
temporary_output = OUTPUT.with_suffix(".tmp.docx")
with ZipFile(REFERENCE, "r") as source_zip, ZipFile(WORKING, "r") as working_zip:
    with ZipFile(temporary_output, "w") as final_zip:
        for info in source_zip.infolist():
            if info.filename in replacement_parts:
                data = working_zip.read(info.filename)
            else:
                data = source_zip.read(info.filename)
            final_zip.writestr(info, data)
temporary_output.replace(OUTPUT)

if file_sha(REFERENCE) != EXPECTED_SHA:
    raise RuntimeError("Reference template was modified during authoring.")

print(OUTPUT)
